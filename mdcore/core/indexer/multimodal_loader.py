from __future__ import annotations
from pathlib import Path

from langchain_core.documents import Document

from mdcore.config.models import VaultConfig
from mdcore.utils.file_utils import vault_relative_path, folder_path_from_relative
from mdcore.utils.logging import get_logger

log = get_logger("indexer.multimodal_loader")


class MultiModalLoader:
    """
    Loads PDF, DOCX, and TXT files into LangChain Documents.

    Contract:
    - Returns a Document with page_content as extracted plain text.
    - Metadata mirrors DocumentLoader: source_file, folder_path, filename.
    - Adds file_type metadata for downstream filtering.
    - Raises ImportError with install hint if required library is missing.
    - Returns Document with empty page_content if extraction yields nothing
      (e.g. scanned PDF). Caller filters on word count.
    """

    def __init__(self, vault_cfg: VaultConfig) -> None:
        self._vault_path = Path(vault_cfg.path).expanduser()

    def load(self, path: Path) -> Document:
        suffix = path.suffix.lower()
        if suffix == ".pdf":
            content = self._load_pdf(path)
        elif suffix == ".docx":
            content = self._load_docx(path)
        elif suffix == ".txt":
            content = path.read_text(encoding="utf-8", errors="ignore")
        else:
            raise ValueError(f"MultiModalLoader: unsupported format {suffix!r}")

        rel = vault_relative_path(path, self._vault_path)
        folder = folder_path_from_relative(rel)

        metadata = {
            "source_file": rel,
            "folder_path": folder,
            "filename": path.name,
            "frontmatter": {},
            "file_type": suffix.lstrip("."),
        }
        log.debug("Loaded %s (%s): %d chars", rel, suffix, len(content))
        return Document(page_content=content, metadata=metadata)

    def _load_pdf(self, path: Path) -> str:
        try:
            from pypdf import PdfReader
        except ImportError:
            raise ImportError(
                "pypdf is required for PDF indexing. "
                "Install with: pip install 'markdowncore-ai[multimodal]'"
            )
        reader = PdfReader(str(path))
        pages = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            if text.strip():
                pages.append(f"## Page {i + 1}\n\n{text}")
        extracted = "\n\n".join(pages)
        if not extracted.strip():
            log.warning("PDF yielded no text (possibly scanned/image-only): %s", path)
        return extracted

    def _load_docx(self, path: Path) -> str:
        try:
            import docx
        except ImportError:
            raise ImportError(
                "python-docx is required for DOCX indexing. "
                "Install with: pip install 'markdowncore-ai[multimodal]'"
            )
        doc = docx.Document(str(path))
        parts = []

        for para in doc.paragraphs:
            if not para.text.strip():
                continue
            style = para.style.name if para.style else ""
            if "Heading 1" in style:
                parts.append(f"# {para.text}")
            elif "Heading 2" in style:
                parts.append(f"## {para.text}")
            elif "Heading 3" in style:
                parts.append(f"### {para.text}")
            else:
                parts.append(para.text)

        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                rows.append("| " + " | ".join(cells) + " |")
            if rows:
                parts.append("\n".join(rows))

        return "\n\n".join(parts)
